"""
Helpers for loading and parsing Word documents (.docx).

This module provides utilities for extracting text from Word document files
(.docx format) and splitting them into chunks. Word processing uses python-docx
which is the standard library for handling Microsoft Word documents.

The splitter follows the same approach as markdown and PDF: it groups text into
chunks until a maximum number of words is reached, with overlap between chunks
to maintain context.
"""

from __future__ import annotations

import os
from typing import List

from ..utils.logging import get_logger

logger = get_logger(__name__)

try:
    from docx import Document
    from docx.document import Document as DocumentType
    from docx.text.paragraph import Paragraph
    from docx.table import Table
    _python_docx_available = True
except ImportError:
    Document = None  # type: ignore
    DocumentType = None  # type: ignore
    Paragraph = None  # type: ignore
    Table = None  # type: ignore
    _python_docx_available = False
    _import_error = ImportError(
        "python-docx is not installed. Install it with: pip install python-docx"
    )


class WordParseError(Exception):
    """Raised when an error occurs while parsing a Word document."""


def load_word_recursive(root: str) -> List[str]:
    """Return a list of Word document file paths under the given root.

    The function descends into all subdirectories and collects files
    ending with ``.docx``. Non-ASCII filenames are supported.
    
    Note: Only .docx format is supported (modern Word format). Older .doc
    files are not supported and would require additional tools like antiword
    or LibreOffice.

    Args:
        root: Base directory to search.
    Returns:
        A list of absolute file paths.
    Raises:
        ValueError: If the root directory doesn't exist or is not a directory.
    """
    if not os.path.exists(root):
        raise ValueError(f"Directory does not exist: {root}")
    if not os.path.isdir(root):
        raise ValueError(f"Path is not a directory: {root}")

    word_files: List[str] = []
    for dirpath, _, filenames in os.walk(root):
        for fname in filenames:
            if fname.lower().endswith(".docx"):
                word_files.append(os.path.join(dirpath, fname))
    return word_files


def extract_text_from_word(docx_path: str) -> str:
    """Extract all text from a Word document (.docx file).

    Uses python-docx to extract text from paragraphs, tables, headers, and
    footers. Handles errors gracefully and logs issues for troubleshooting.
    
    The extraction includes:
    - All paragraphs from the main document body
    - Text from tables (cells are separated by newlines)
    - Headers and footers (if present)
    
    Note: Images, embedded objects, and complex formatting are not extracted.
    Only text content is extracted.

    Args:
        docx_path: Path to the .docx file.
    Returns:
        Extracted text content as a single string.
    Raises:
        WordParseError: If the document cannot be read or parsed.
        ImportError: If python-docx is not installed.
    """
    if not _python_docx_available or Document is None:
        raise ImportError(
            "python-docx is not installed. Install it with: pip install python-docx"
        )

    if not os.path.exists(docx_path):
        raise WordParseError(f"Word document does not exist: {docx_path}")

    try:
        text_parts: List[str] = []
        
        # Open the document
        doc: DocumentType = Document(docx_path)
        
        # Extract text from paragraphs in the main document body
        for para in doc.paragraphs:
            para_text = para.text.strip()
            if para_text:
                text_parts.append(para_text)
        
        # Extract text from tables
        for table in doc.tables:
            table_text_parts: List[str] = []
            for row in table.rows:
                row_text_parts: List[str] = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text_parts.append(cell_text)
                if row_text_parts:
                    # Join cells in a row with tab separator
                    table_text_parts.append(" | ".join(row_text_parts))
            if table_text_parts:
                # Add table text with separator
                text_parts.append("\n".join(table_text_parts))
        
        # Extract text from headers and footers (if present)
        # Note: Headers/footers are typically less important for indexing
        # but can provide context. We extract them but mark them clearly.
        try:
            for section in doc.sections:
                # Header
                try:
                    header = section.header
                    for para in header.paragraphs:
                        header_text = para.text.strip()
                        if header_text:
                            text_parts.append(f"[Header] {header_text}")
                except Exception:
                    # Some documents may not have headers
                    pass
                
                # Footer
                try:
                    footer = section.footer
                    for para in footer.paragraphs:
                        footer_text = para.text.strip()
                        if footer_text:
                            text_parts.append(f"[Footer] {footer_text}")
                except Exception:
                    # Some documents may not have footers
                    pass
        except Exception:
            # Headers/footers extraction is optional, continue if it fails
            logger.debug("Could not extract headers/footers from Word document")

        if not text_parts:
            logger.warning(f"No text extracted from Word document: {docx_path}")
            return ""

        # Join all text parts with double newline (similar to markdown paragraph separation)
        full_text = "\n\n".join(text_parts)
        return full_text

    except Exception as exc:
        logger.error(f"Failed to parse Word document {docx_path}: {exc}")
        raise WordParseError(f"Failed to parse Word document {docx_path}: {exc}") from exc


def split_word_text(text: str, max_words: int = 300, overlap: int = 50) -> List[str]:
    """Split Word document text into chunks of roughly ``max_words`` words.

    The algorithm splits the text into paragraphs on blank lines and
    accumulates paragraphs until the running total exceeds ``max_words``.
    When a chunk is emitted, the last ``overlap`` words of the previous
    chunk are prepended to the next chunk to provide context overlap.

    Args:
        text: Word document text as a single string.
        max_words: Target number of words per chunk.
        overlap: Number of words to repeat between consecutive chunks.
    Returns:
        A list of text chunks.
    """
    if not text.strip():
        return []

    # Split into paragraphs (similar to markdown and PDF processing)
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks: List[str] = []
    current_words: List[str] = []

    for para in paragraphs:
        words = para.split()
        if not words:
            continue
        # If adding this paragraph would exceed the maximum, emit a chunk
        if len(current_words) + len(words) > max_words:
            if current_words:
                chunks.append(" ".join(current_words))
                # Prepare overlap for next chunk
                current_words = current_words[-overlap:] + words
            else:
                # Single paragraph longer than max_words - split it
                chunks.append(" ".join(words))
                current_words = []
        else:
            current_words.extend(words)

    # Emit remaining words
    if current_words:
        chunks.append(" ".join(current_words))

    return chunks

